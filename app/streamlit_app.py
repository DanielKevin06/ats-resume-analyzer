import streamlit as st
import requests
import json
import io
import re
import random
from datetime import datetime
from collections import Counter

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

import plotly.graph_objects as go
import plotly.express as px
import streamlit.components.v1 as components

import os
API_URL = os.getenv("API_URL", "http://127.0.0.1:8000")

st.set_page_config(page_title="Resume ATS Analyzer", page_icon="📄", layout="wide")

# ============================================================
# OPTIONAL IMPORTS for JD FILE EXTRACTION
# ============================================================
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    from docx import Document
except Exception:
    Document = None


# ============================================================
# STYLES (LIGHT UI + ✅ GREY BUTTONS + ✅ FORCE WHITE TEXT)
# ============================================================
st.markdown(
    """
<style>
html, body, .stApp { background: #ffffff !important; color: #111827 !important; }
[data-testid="stAppViewContainer"] { background: #ffffff !important; }
[data-testid="stHeader"] { background: rgba(255,255,255,0) !important; }

h1, h2, h3, h4, h5, h6 { color: #111827 !important; font-weight: 800 !important; }
p, label, span, div { color: #111827; }

.stTextArea textarea, .stTextInput input {
    background: #ffffff !important;
    color: #111827 !important;
    border: 1px solid #e5e7eb !important;
    border-radius: 12px !important;
}

/* Uploader dropzone */
[data-testid="stFileUploaderDropzone"] {
    background: #f9fafb !important;
    border: 1px dashed #d1d5db !important;
    border-radius: 12px !important;
}
[data-testid="stFileUploaderDropzoneInstructions"] * { color: #111827 !important; }

/* ============================================================
   ✅ ALL STREAMLIT BUTTONS = GREY + FORCE WHITE TEXT
   Covers st.button + st.download_button across Streamlit versions
   ============================================================ */
.stButton > button,
div.stButton > button,
[data-testid="stDownloadButton"] > button,
div[data-testid="stDownloadButton"] > button {
    background: #6b7280 !important;   /* grey */
    color: #ffffff !important;        /* white text */
    height: 52px !important;
    font-size: 16px !important;
    font-weight: 800 !important;
    border-radius: 12px !important;
    border: none !important;
    box-shadow: 0 4px 12px rgba(0,0,0,0.12) !important;
}

/* ✅ FORCE WHITE INSIDE BUTTON (text/icons) */
.stButton > button * ,
div.stButton > button * ,
[data-testid="stDownloadButton"] > button * ,
div[data-testid="stDownloadButton"] > button * {
    color: #ffffff !important;
    fill: #ffffff !important;
    stroke: #ffffff !important;
}

/* Hover/Active */
.stButton > button:hover,
div.stButton > button:hover,
[data-testid="stDownloadButton"] > button:hover,
div[data-testid="stDownloadButton"] > button:hover {
    background: #4b5563 !important;
}

.stButton > button:active,
div.stButton > button:active,
[data-testid="stDownloadButton"] > button:active,
div[data-testid="stDownloadButton"] > button:active {
    background: #374151 !important;
}

/* Cards */
.card {
    background: #ffffff !important;
    padding: 18px !important;
    border-radius: 14px !important;
    border: 1px solid #e5e7eb !important;
    box-shadow: 0 6px 18px rgba(0,0,0,0.06) !important;
    overflow: visible !important;
}

/* FIX File Uploader cancel button (keep red + proper X icon) */
[data-testid="stFileUploader"] { overflow: visible !important; }

[data-testid="stFileUploaderFile"]{
    display:flex !important;
    align-items:center !important;
    justify-content:space-between !important;
    overflow:visible !important;
}

/* red button */
[data-testid="stFileUploaderFile"] button{
    position:relative !important;
    margin-left:12px !important;
    padding:6px 10px !important;
    background:#ef4444 !important;
    color:#ffffff !important;
    border-radius:10px !important;
    border:none !important;
    cursor:pointer !important;
}

/* ✅ IMPORTANT: don’t force fill on everything (prevents white square) */
[data-testid="stFileUploaderFile"] button svg{
    width:18px !important;
    height:18px !important;
}

/* make the X visible */
[data-testid="stFileUploaderFile"] button svg path{
  stroke:#ffffff !important;
  stroke-width:2.4 !important;   /* thicker X */
  fill:none !important;
}


/* if svg uses rect/circle, keep them transparent */
[data-testid="stFileUploaderFile"] button svg rect,
[data-testid="stFileUploaderFile"] button svg circle{
    fill:none !important;
    stroke:#ffffff !important;
}

[data-testid="stFileUploaderFile"] button:hover{ background:#dc2626 !important; }


/* Chips */
.chip {
    display: inline-block;
    padding: 6px 10px;
    border-radius: 999px;
    font-size: 13px;
    border: 1px solid #e5e7eb;
    margin: 0;
}
.chip-wrap{
    display: flex;
    flex-wrap: wrap;
    gap: 10px;
    align-items: flex-start;
}
.ok   { background:#ecfdf5 !important; color:#065f46 !important; border-color:#a7f3d0 !important; }
.miss { background:#fff1f2 !important; color:#9f1239 !important; border-color:#fecdd3 !important; }

/* Sidebar history cards */
.sidebar-card{
    background:#ffffff;
    border:1px solid #e5e7eb;
    border-radius:12px;
    padding:12px;
    margin-bottom:10px;
    box-shadow: 0 4px 14px rgba(0,0,0,0.04);
}
.sidebar-muted{ color:#6b7280 !important; font-size:12px; }

hr { border: none; border-top: 1px solid #e5e7eb; margin: 12px 0; }

/* Suggestion rows */
.sugg-text{
    padding:10px 12px;
    border:1px solid #e5e7eb;
    border-radius:12px;
    background:#f9fafb;
}
</style>
""",
    unsafe_allow_html=True,
)

# ============================================================
# HELPERS
# ============================================================
def now_label():
    return datetime.now().strftime("%Y-%m-%d %H:%M")


def normalize_similarity(raw_sim):
    try:
        raw_sim = float(raw_sim)
    except Exception:
        raw_sim = 0.0
    sim_norm = raw_sim / 100.0 if raw_sim > 1 else raw_sim
    sim_norm = max(0.0, min(1.0, sim_norm))
    return sim_norm


def render_chips(chips, css_class="ok", limit=150):
    if not chips:
        st.write("None")
        return
    chips = chips[:limit]
    html = '<div class="chip-wrap">' + "".join(
        [f'<span class="chip {css_class}">{c}</span>' for c in chips]
    ) + "</div>"
    st.markdown(html, unsafe_allow_html=True)


def make_pdf_report(result: dict, resume_name: str) -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 50

    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y, "Resume ATS Analyzer Report")
    y -= 24

    c.setFont("Helvetica", 11)
    c.drawString(50, y, f"Resume: {resume_name}")
    y -= 16
    c.drawString(50, y, f"ATS Score: {result.get('total_score', 0)}/100")
    y -= 16

    sim_norm = normalize_similarity(result.get("similarity", 0.0))
    c.drawString(50, y, f"Similarity: {round(sim_norm*100, 2)}%")
    y -= 24

    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "Matched Keywords (Top 25)")
    y -= 14
    c.setFont("Helvetica", 9)
    c.drawString(50, y, ", ".join(result.get("keywords_matched", [])[:25])[:120])
    y -= 22

    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "Missing Keywords (Top 25)")
    y -= 14
    c.setFont("Helvetica", 9)
    c.drawString(50, y, ", ".join(result.get("keywords_missing", [])[:25])[:120])
    y -= 22

    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "Action Items")
    y -= 14
    c.setFont("Helvetica", 9)
    for item in (result.get("action_items", []) or [])[:6]:
        c.drawString(55, y, f"- {item}"[:115])
        y -= 12

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer.read()


def ats_gauge(score: int):
    fig = go.Figure(
        go.Indicator(
            mode="gauge+number",
            value=score,
            number={"font": {"size": 48}},
            title={"text": "ATS Score"},
            gauge={
                "axis": {"range": [0, 100]},
                "bar": {"color": "#111827"},
                "steps": [
                    {"range": [0, 40], "color": "#fee2e2"},
                    {"range": [40, 70], "color": "#fde68a"},
                    {"range": [70, 100], "color": "#bbf7d0"},
                ],
            },
        )
    )
    fig.update_layout(height=320, margin=dict(l=20, r=20, t=40, b=20))
    return fig


def extract_text_from_upload(file) -> str:
    """Supports: .txt, .pdf, .docx"""
    if file is None:
        return ""

    name = (file.name or "").lower()

    if name.endswith(".txt"):
        try:
            return file.getvalue().decode("utf-8", errors="ignore")
        except Exception:
            return file.getvalue().decode(errors="ignore")

    if name.endswith(".pdf"):
        if pdfplumber is None:
            st.error("PDF reading needs `pdfplumber`. Install: pip install pdfplumber")
            return ""
        try:
            text_parts = []
            with pdfplumber.open(io.BytesIO(file.getvalue())) as pdf:
                for page in pdf.pages:
                    text_parts.append(page.extract_text() or "")
            return "\n".join(text_parts).strip()
        except Exception as e:
            st.error(f"Could not read PDF: {e}")
            return ""

    if name.endswith(".docx"):
        if Document is None:
            st.error("DOCX reading needs `python-docx`. Install: pip install python-docx")
            return ""
        try:
            doc = Document(io.BytesIO(file.getvalue()))
            return "\n".join([p.text for p in doc.paragraphs]).strip()
        except Exception as e:
            st.error(f"Could not read DOCX: {e}")
            return ""

    st.error("Unsupported file type. Please upload PDF, DOCX, or TXT.")
    return ""


def jd_top_terms(text: str, top_n: int = 20):
    if not text or not text.strip():
        return [], []

    words = re.findall(r"[a-zA-Z]{3,}", text.lower())

    stop = {
        "the","and","for","with","that","this","you","your","are","will","from","have",
        "has","was","were","their","they","them","our","but","not","can","able","all",
        "any","into","over","more","less","than","also","such","use","using","used",
        "job","role","work","working","years","year","experience","skills","skill",
        "responsibilities","responsibility"
    }
    words = [w for w in words if w not in stop]

    if not words:
        return [], []

    counts = Counter(words).most_common(top_n)
    terms = [t for t, _ in counts]
    freqs = [c for _, c in counts]
    return terms, freqs


# ============================================================
# ✅ COPY BUTTON (works, no Streamlit key errors)
# ============================================================
def copy_to_clipboard_button(text: str, unique_id: str, label: str = "📋 Copy"):
    safe_text = text.replace("\\", "\\\\").replace("`", "\\`").replace("${", "\\${")

    html = f"""
    <button id="{unique_id}"
        style="
            width:100%;
            padding:10px 12px;
            border-radius:12px;
            border:1px solid #e5e7eb;
            background:#ffffff;
            cursor:pointer;
            font-weight:700;
        "
        onclick="
            navigator.clipboard.writeText(`{safe_text}`);
            const btn = document.getElementById('{unique_id}');
            btn.innerText='✅ Copied';
            setTimeout(()=>{{ btn.innerText='{label}'; }}, 1500);
        "
    >{label}</button>
    """
    components.html(html, height=52)


# ============================================================
# ✅ AUTO RESUME BULLET SUGGESTIONS (JD + missing)
# ============================================================
def infer_role_hint(jd_text: str) -> str:
    jd = (jd_text or "").lower()
    if any(k in jd for k in ["soc", "siem", "incident response", "splunk", "cyber", "firewall", "vulnerability"]):
        return "security"
    if any(k in jd for k in ["react", "frontend", "ui", "css", "javascript", "typescript"]):
        return "frontend"
    if any(k in jd for k in ["backend", "spring", "microservices", "api", "java", "node", "fastapi", "django"]):
        return "backend"
    if any(k in jd for k in ["aws", "azure", "gcp", "docker", "kubernetes", "devops", "ci/cd"]):
        return "devops"
    if any(k in jd for k in ["machine learning", "ml", "model", "pandas", "python", "data science"]):
        return "data"
    return "general"


def generate_resume_suggestions(missing_keywords, jd_text: str, limit: int = 8):
    role = infer_role_hint(jd_text)

    templates_general = [
        "Implemented {k} solutions to improve performance, reliability, and scalability.",
        "Built and optimized {k} workflows aligned with business requirements.",
        "Developed production-ready {k} components with strong documentation and testing.",
        "Collaborated with cross-functional teams to deliver {k}-driven outcomes on time.",
        "Improved operational efficiency by integrating {k} into existing processes.",
    ]

    templates_security = [
        "Supported incident response workflows using {k} to improve threat detection and triage.",
        "Implemented security controls and monitoring with {k} to reduce risk exposure.",
        "Performed vulnerability assessment and remediation activities using {k}.",
        "Built security reporting and compliance evidence using {k} and documented procedures.",
    ]

    templates_devops = [
        "Automated deployment workflows using {k} to improve release speed and stability.",
        "Built CI/CD pipelines leveraging {k} with rollback and monitoring best practices.",
        "Containerized applications using {k} and improved environment consistency.",
        "Improved observability and uptime using {k}-based monitoring and alerting.",
    ]

    templates_data = [
        "Built {k}-driven analytics pipelines to improve insight generation and reporting.",
        "Developed data workflows using {k} to improve quality, reliability, and performance.",
        "Applied {k} methods to analyze trends and communicate results to stakeholders.",
        "Implemented {k} solutions for scalable data processing and experimentation.",
    ]

    templates_frontend = [
        "Built responsive UI features using {k} with a focus on performance and accessibility.",
        "Implemented reusable components using {k} and improved UX consistency.",
        "Integrated APIs into the UI using {k} and improved error handling and loading states.",
        "Improved frontend performance using {k} best practices and optimization.",
    ]

    templates_backend = [
        "Designed and implemented backend services using {k} with clean architecture practices.",
        "Built REST APIs using {k} with authentication, validation, and logging.",
        "Optimized backend performance using {k} and improved scalability under load.",
        "Implemented database and integration workflows using {k} with strong error handling.",
    ]

    role_map = {
        "security": templates_security,
        "devops": templates_devops,
        "data": templates_data,
        "frontend": templates_frontend,
        "backend": templates_backend,
        "general": templates_general,
    }
    templates = role_map.get(role, templates_general)

    cleaned = []
    seen = set()
    for kw in (missing_keywords or []):
        k = str(kw).strip()
        if not k:
            continue
        low = k.lower()
        if low in seen:
            continue
        seen.add(low)
        cleaned.append(k)

    bullets = []
    for k in cleaned[:max(limit, 1)]:
        t = random.choice(templates)
        bullets.append(f"• {t.format(k=k.title())}")

    return bullets


# ============================================================
# SESSION STATE INIT
# ============================================================
st.session_state.setdefault("jd_text", "")
st.session_state.setdefault("last_result", None)
st.session_state.setdefault("last_resume_name", "resume.pdf")
st.session_state.setdefault("last_jd_text", "")
st.session_state.setdefault("history", [])
st.session_state.setdefault("pinned_best_id", None)


def add_run_to_history(resume_name: str, jd_text: str, result: dict):
    sim_norm = normalize_similarity(result.get("similarity", 0.0))
    run = {
        "id": f"{resume_name}-{datetime.now().timestamp()}",
        "resume": resume_name,
        "ats": int(result.get("total_score", 0)),
        "sim": round(sim_norm * 100, 1),
        "time": now_label(),
        "result": result,
        "jd_text": jd_text,
    }
    st.session_state["history"].append(run)

    best = max(st.session_state["history"], key=lambda r: (r["ats"], r["sim"]), default=None)
    st.session_state["pinned_best_id"] = best["id"] if best else None


# ============================================================
# SIDEBAR
# ============================================================
with st.sidebar:
    st.markdown("## 📌 History")
    st.write("Run an analysis to save history here.")

    if st.button("🧹 Clear History", use_container_width=True):
        st.session_state["history"] = []
        st.session_state["pinned_best_id"] = None
        st.success("History cleared ✅")
        st.rerun()

    if st.button("🗑️ Reset All", use_container_width=True):
        st.session_state["history"] = []
        st.session_state["pinned_best_id"] = None
        st.session_state["last_result"] = None
        st.session_state["last_resume_name"] = "resume.pdf"
        st.session_state["jd_text"] = ""
        st.session_state["last_jd_text"] = ""
        st.session_state.pop("compare_A", None)
        st.session_state.pop("compare_B", None)
        st.success("Reset done ✅")
        st.rerun()

    st.markdown("---")

    history = st.session_state.get("history", [])

    st.markdown("### 🏆 Best Run (Pinned)")
    best_id = st.session_state.get("pinned_best_id", None)
    best_run = next((r for r in history if r.get("id") == best_id), None)

    if best_run:
        st.markdown(
            f"""
<div class="sidebar-card">
  <div><b>{best_run.get("resume","resume.pdf")}</b></div>
  <div class="sidebar-muted">ATS: {best_run.get("ats","-")}/100 • Similarity: {best_run.get("sim","-")}%</div>
  <div class="sidebar-muted">{best_run.get("time","")}</div>
</div>
""",
            unsafe_allow_html=True,
        )

        if st.button("📌 Load Best Run", use_container_width=True):
            st.session_state["last_result"] = best_run.get("result")
            st.session_state["last_resume_name"] = best_run.get("resume", "resume.pdf")
            st.session_state["last_jd_text"] = best_run.get("jd_text", "")
            st.success("Loaded best run ✅")
            st.rerun()
    else:
        st.caption("No runs yet.")

    st.markdown("---")
    st.markdown("### 🧾 Runs")

    if not history:
        st.markdown(
            '<div class="sidebar-card"><span class="sidebar-muted">No history yet.</span></div>',
            unsafe_allow_html=True,
        )
    else:
        for item in history[::-1][:8]:
            st.markdown(
                f"""
<div class="sidebar-card">
  <div><b>{item.get("resume","resume.pdf")}</b></div>
  <div class="sidebar-muted">ATS: {item.get("ats","-")}/100 • Similarity: {item.get("sim","-")}%</div>
  <div class="sidebar-muted">{item.get("time","")}</div>
</div>
""",
                unsafe_allow_html=True,
            )

    st.markdown("---")
    st.markdown("### 🔁 Compare Runs")

    if len(history) < 2:
        st.caption("Need at least 2 runs to compare.")
    else:
        options = [(r["id"], f"{r['resume']} | ATS {r['ats']}/100 | {r['time']}") for r in history[::-1]]

        run_a = st.selectbox("Select Run A", options, format_func=lambda x: x[1], key="run_a")[0]
        run_b = st.selectbox("Select Run B", options, format_func=lambda x: x[1], key="run_b")[0]

        if st.button("⚡ Compare Selected Runs", use_container_width=True):
            A = next((r for r in history if r["id"] == run_a), None)
            B = next((r for r in history if r["id"] == run_b), None)

            if not A or not B:
                st.error("Could not load selected runs.")
            else:
                st.session_state["compare_A"] = A
                st.session_state["compare_B"] = B
                st.success("Comparison loaded ✅")
                st.rerun()


# ============================================================
# MAIN UI
# ============================================================
st.title("Resume ATS Analyzer")
st.caption("Upload resume + paste or upload Job Description to evaluate ATS score")

col1, col2 = st.columns([1, 1], gap="large")

with col1:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Upload & Analyze")

    resume_file = st.file_uploader("Upload Resume (PDF/DOCX)", type=["pdf", "docx"])

    jd_file = st.file_uploader("Upload Job Description (PDF/DOCX/TXT) (Optional)", type=["pdf", "docx", "txt"])
    if jd_file is not None:
        jd_text_from_file = extract_text_from_upload(jd_file)
        if jd_text_from_file.strip():
            st.session_state["jd_text"] = jd_text_from_file
            st.success("JD text loaded from file. You can edit it below.")

    jd = st.text_area(
        "Paste Job Description",
        height=220,
        placeholder="Paste full JD here...",
        key="jd_text",
    )

    analyze_btn = st.button("🔍 Analyze", use_container_width=True)
    st.markdown("</div>", unsafe_allow_html=True)

with col2:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Results")
    st.write("Run an analysis to view results")
    if st.session_state.get("last_result"):
        st.info("Tip: Your last analysis is saved. Refresh won’t lose results.")
    st.markdown("</div>", unsafe_allow_html=True)


# ============================================================
# ANALYZE
# ============================================================
if analyze_btn:
    if not resume_file:
        st.error("Please upload a resume first.")
        st.stop()

    jd_to_send = (jd or "").strip()
    if not jd_to_send:
        st.error("Please paste or upload a Job Description.")
        st.stop()

    with st.spinner("Analyzing..."):
        files = {"resume": (resume_file.name, resume_file.getvalue(), resume_file.type)}
        data = {"jd": jd_to_send}
        resp = requests.post(f"{API_URL}/analyze", files=files, data=data, timeout=120)

        try:
            result = resp.json()
        except Exception:
            st.error("Backend did not return JSON. Check FastAPI logs.")
            st.stop()

    st.session_state["last_result"] = result
    st.session_state["last_resume_name"] = resume_file.name
    st.session_state["last_jd_text"] = jd_to_send

    if "error" not in result:
        add_run_to_history(resume_file.name, jd_to_send, result)


# ============================================================
# RENDER RESULTS
# ============================================================
result = st.session_state.get("last_result", None)
resume_name = st.session_state.get("last_resume_name", "resume.pdf")
jd_current = st.session_state.get("last_jd_text", "") or st.session_state.get("jd_text", "")

if result:
    if "error" in result:
        st.error(result["error"])
        st.stop()

    score = int(result.get("total_score", 0))
    matched = result.get("keywords_matched", []) or []
    missing = result.get("keywords_missing", []) or []
    tips = result.get("action_items", []) or []

    sim_norm = normalize_similarity(result.get("similarity", 0.0))
    sim_pct = round(sim_norm * 100, 1)

    if score >= 80:
        st.success("🔥 Excellent ATS Score — your resume is strong for this JD!")
    elif score >= 60:
        st.warning("⚠️ Medium ATS Score — some improvements will help.")
    else:
        st.error("❌ Low ATS Score — needs significant keyword & structure improvements.")

    st.subheader("Score Overview")
    st.plotly_chart(ats_gauge(score), use_container_width=True)

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("ATS Score", f"{score}/100")
    m2.metric("Similarity", f"{sim_pct}%")
    m3.metric("Keywords Matched", str(len(matched)))
    m4.metric("Keywords Missing", str(len(missing)))

    st.subheader("Keyword Coverage")
    fig_pie = px.pie(values=[len(matched), len(missing)], names=["Matched", "Missing"], title="Matched vs Missing Keywords")
    st.plotly_chart(fig_pie, use_container_width=True)

    st.subheader("Similarity Score")
    st.metric("Similarity", f"{sim_pct}%")
    st.progress(sim_norm)

    st.subheader("📈 JD Keyword Frequency (Top 20)")
    terms, counts = jd_top_terms(jd_current, top_n=20)
    if not terms:
        st.info("Not enough JD text to compute keyword frequency (or everything was filtered as common words).")
    else:
        freq_df = {"Keyword": terms, "Count": counts}
        fig_bar = px.bar(freq_df, x="Keyword", y="Count", title="Top JD terms")
        st.plotly_chart(fig_bar, use_container_width=True)

    st.subheader("Keyword Search")
    search = st.text_input("Type to filter keywords (matched + missing)", placeholder="e.g., python, aws, tableau")

    if search:
        s = search.lower().strip()
        matched_filtered = [k for k in matched if s in k.lower()]
        missing_filtered = [k for k in missing if s in k.lower()]
    else:
        matched_filtered = matched
        missing_filtered = missing

    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Matched Keywords")
    if matched_filtered:
        render_chips(matched_filtered, "ok", limit=250)
    else:
        st.write("No matched keywords detected (or none match your filter).")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Missing Keywords")
    if missing_filtered:
        st.markdown("**Top Missing Keywords**")
        render_chips(missing_filtered, "miss", limit=40)
        with st.expander("Show all missing keywords"):
            render_chips(missing_filtered, "miss", limit=500)
    else:
        st.write("No missing keywords detected (or none match your filter).")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Improvement Tips")
    if tips:
        for tip in tips[:10]:
            st.write("•", tip)
    else:
        st.write("No action items provided by backend. (Optional: add action_items in API response.)")
    st.markdown("</div>", unsafe_allow_html=True)

    # ✅ AUTO BULLET SUGGESTIONS + COPY
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("✨ Suggested Resume Bullet Points (Auto-generated)")
    st.caption("These suggestions are generated dynamically using the JD and your missing keywords.")

    if missing_filtered:
        suggestions = generate_resume_suggestions(missing_filtered, jd_current, limit=8)

        for i, s_text in enumerate(suggestions, start=1):
            c_left, c_right = st.columns([4, 1])
            with c_left:
                st.markdown(f"<div class='sugg-text'>{s_text}</div>", unsafe_allow_html=True)
            with c_right:
                copy_to_clipboard_button(s_text, unique_id=f"copy_btn_{i}", label="📋 Copy")

        st.markdown(" ")
        all_text = "\n".join(suggestions).encode("utf-8")
        st.download_button(
            "⬇️ Download Suggestions (.txt)",
            all_text,
            "resume_suggestions.txt",
            "text/plain",
            use_container_width=True,
        )
    else:
        st.info("No missing keywords found — suggestions not needed. Try a different JD or resume.")

    st.markdown("</div>", unsafe_allow_html=True)

    # Download
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Download Report")

    json_bytes = json.dumps(result, indent=2).encode("utf-8")
    st.download_button("⬇️ Download JSON", json_bytes, "ats_report.json", "application/json", use_container_width=True)

    pdf_bytes = make_pdf_report(result, resume_name)
    st.download_button("⬇️ Download PDF", pdf_bytes, "ats_report.pdf", "application/pdf", use_container_width=True)
    st.markdown("</div>", unsafe_allow_html=True)

    with st.expander("Debug JSON"):
        st.json(result)


# ============================================================
# COMPARISON VIEW (if loaded)
# ============================================================
A = st.session_state.get("compare_A")
B = st.session_state.get("compare_B")
if A and B:
    st.markdown("---")
    st.subheader("🔁 Run Comparison")

    c1, c2, c3 = st.columns([1, 1, 1])
    c1.metric("Run A ATS", f"{A['ats']}/100")
    c2.metric("Run B ATS", f"{B['ats']}/100")
    c3.metric("Δ ATS (B-A)", f"{B['ats'] - A['ats']}")

    A_miss = set((A["result"].get("keywords_missing") or []))
    B_miss = set((B["result"].get("keywords_missing") or []))

    improved = sorted(list(A_miss - B_miss))
    regressed = sorted(list(B_miss - A_miss))

    st.markdown("**✅ Improved (no longer missing in Run B):**")
    if improved:
        render_chips(improved, "ok", limit=60)
    else:
        st.caption("None")

    st.markdown("**⚠️ New missing in Run B:**")
    if regressed:
        render_chips(regressed, "miss", limit=60)
    else:
        st.caption("None")
